# -*- coding: utf-8 -*-
"""Word 简历排版解析 → 定制 CSS（完美还原排版）测试。

覆盖：字体/字号/加粗/颜色/对齐/缩进/间距/行距 提取 + CSS 还原 + executor 工具接入。
"""
import os
import sys

_SRC = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src")
if _SRC not in sys.path:
    sys.path.insert(0, _SRC)

import pytest

docx = pytest.importorskip("docx")
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from resume_product.render.docx_import import parse_docx, import_docx_resume


def _set_font(run, name, size_pt, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _make_resume_docx(path):
    """生成带排版的测试简历 docx。"""
    doc = Document()
    # 姓名：黑体 22pt 居中加粗
    p = doc.add_paragraph()
    _set_font(p.add_run("张 三"), "黑体", 22, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    # 联系方式：宋体 10.5pt 居中
    p = doc.add_paragraph()
    _set_font(p.add_run("138-0000-0000 | zhangsan@email.com"), "宋体", 10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 章节：黑体 14pt 加粗，段前 10 段后 6
    p = doc.add_paragraph()
    _set_font(p.add_run("工作经历"), "黑体", 14, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    # 正文：宋体 11pt 首行缩进 22pt 1.5 倍行距
    p = doc.add_paragraph()
    _set_font(p.add_run("在某公司担任工程师，负责后端开发与系统设计。"), "宋体", 11)
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    # 混合 run：加粗 + 红色
    p = doc.add_paragraph()
    _set_font(p.add_run("关键成果："), "宋体", 11, bold=True)
    r2 = p.add_run("性能提升 30%")
    _set_font(r2, "宋体", 11, color=(0xC0, 0x00, 0x00))
    doc.save(path)


@pytest.fixture
def resume_docx(tmp_path):
    p = str(tmp_path / "resume.docx")
    _make_resume_docx(p)
    return p


def test_parse_docx_extracts_paragraphs(resume_docx):
    r = parse_docx(resume_docx)
    paras = r["paragraphs"]
    assert len(paras) == 5, "应解析出 5 个非空段落"

    # 姓名段
    name = paras[0]
    assert name["text"] == "张 三"
    assert name["alignment"] == "center"
    assert name["space_after_pt"] == 12.0
    assert name["runs"][0]["font_name"] == "黑体"
    assert name["runs"][0]["size_pt"] == 22.0
    assert name["runs"][0]["bold"] is True

    # 章节段
    sec = paras[2]
    assert sec["space_before_pt"] == 10.0
    assert sec["space_after_pt"] == 6.0
    assert sec["runs"][0]["bold"] is True

    # 正文段
    body = paras[3]
    assert body["first_line_indent_pt"] == 22.0
    assert body["line_spacing"] == "line-height:1.5"

    # 混合 run 段（颜色）
    mix = paras[4]
    assert mix["runs"][0]["bold"] is True
    assert mix["runs"][1]["color"] == "#C00000"


def test_import_docx_resume_returns_css_html(resume_docx):
    r = import_docx_resume(resume_docx)
    assert r["meta"]["paragraphs"] == 5
    assert r["meta"]["fonts"] == ["宋体", "黑体"]
    # CSS 还原排版
    assert "font-family:'黑体'" in r["css"], "CSS 应含黑体字体"
    assert "font-size:22.0pt" in r["css"], "CSS 应含姓名 22pt 字号"
    assert "text-align:center" in r["css"], "CSS 应含居中"
    assert "text-indent:22.0pt" in r["css"], "CSS 应含首行缩进"
    assert "line-height:1.5" in r["css"], "CSS 应含 1.5 倍行距"
    assert "margin-top:10.0pt" in r["css"], "CSS 应含段前间距"
    # HTML 还原 run 级差异
    assert "font-weight:bold" in r["html"], "HTML 应含加粗"
    assert "color:#C00000" in r["html"], "HTML 应含红色"
    # 去重：5 段 5 种排版
    assert r["meta"]["distinct_styles"] == 5


def test_executor_import_resume_docx(resume_docx):
    import json
    from resume_product.executor import execute
    out = execute("import_resume_docx", {"docx_path": resume_docx})
    d = json.loads(out)
    assert d["ok"] is True
    assert "css" in d and "html" in d and "meta" in d
    assert "font-family:'黑体'" in d["css"]


def test_import_docx_with_table(tmp_path):
    """Word 表格 → HTML <table>（复杂排版元素还原）。"""
    from docx.oxml.ns import qn
    doc = Document()
    p = doc.add_paragraph()
    _set_font(p.add_run("技能表"), "黑体", 14, bold=True)
    data = [["技能", "水平"], ["Python", "精通"], ["SQL", "熟练"]]
    t = doc.add_table(rows=len(data), cols=2)
    t.style = "Table Grid"
    for i, row in enumerate(data):
        for j, v in enumerate(row):
            t.cell(i, j).text = v
    path = str(tmp_path / "table.docx")
    doc.save(path)

    r = import_docx_resume(path)
    assert r["meta"]["tables"] == 1
    assert "<table" in r["html"], "HTML 应含表格"
    assert "Python" in r["html"] and "精通" in r["html"], "应含单元格内容"
    assert "border-collapse" in r["html"], "表格应含边框样式"


def test_import_docx_with_columns(tmp_path):
    """Word 分栏 → CSS column-count（双栏简历还原）。"""
    from docx.oxml.ns import qn
    doc = Document()
    _set_font(doc.add_paragraph().add_run("双栏简历标题"), "黑体", 16, bold=True)
    _set_font(doc.add_paragraph().add_run("左栏内容"), "宋体", 11)
    sectPr = doc.sections[0]._sectPr
    cols = sectPr.find(qn("w:cols"))
    cols.set(qn("w:num"), "2")
    path = str(tmp_path / "columns.docx")
    doc.save(path)

    r = import_docx_resume(path)
    assert r["meta"]["columns"] == 2
    assert "column-count:2" in r["css"], "CSS 应含分栏声明"


def test_import_docx_with_lists(tmp_path):
    """Word 项目符号/编号列表 → 列表符号还原（真实简历常见）。"""
    doc = Document()
    _set_font(doc.add_paragraph(style="List Bullet").add_run("清华大学 硕士"), "宋体", 11)
    _set_font(doc.add_paragraph(style="List Bullet").add_run("北京大学 学士"), "宋体", 11)
    _set_font(doc.add_paragraph(style="List Number").add_run("字节跳动 算法工程师"), "宋体", 11)
    _set_font(doc.add_paragraph(style="List Number").add_run("美团 后端工程师"), "宋体", 11)
    path = str(tmp_path / "lists.docx")
    doc.save(path)

    r = import_docx_resume(path)
    assert "• " in r["html"], "项目符号应还原为 •"
    assert "1. " in r["html"] and "2. " in r["html"], "编号列表应还原为 1. 2."
    # bullet 应出现两次（两个 bullet 项）
    assert r["html"].count("• ") >= 2


def test_import_docx_with_merged_cells_and_header(tmp_path):
    """合并单元格（colspan）+ 页眉页脚还原。"""
    doc = Document()
    # 页眉
    doc.sections[0].header.paragraphs[0].text = "李明 · 138-0000-0000"
    # 正文 + 合并单元格表格
    doc.add_paragraph().add_run("个人简历")
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "技能总览"
    t.cell(0, 0).merge(t.cell(0, 1))
    t.cell(1, 0).text = "Python"
    t.cell(1, 1).text = "精通"
    path = str(tmp_path / "merge.docx")
    doc.save(path)

    r = import_docx_resume(path)
    assert r["meta"]["headers"] == 1
    assert "李明 · 138-0000-0000" in r["html"], "页眉应还原"
    assert 'colspan="2"' in r["html"], "合并单元格应还原为 colspan"
    assert "技能总览" in r["html"]
