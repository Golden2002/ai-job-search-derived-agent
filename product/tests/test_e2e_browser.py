# -*- coding: utf-8 -*-
"""网页产品浏览器端 E2E（Playwright）—— 场景卡片加载 / 对话收集 / 生成预览 / 导出 / 导入。

覆盖 product/web/index.html（对话式简历制作前端）与 product/src/resume_product/api.py
（Flask 后端）的真实 HTTP 契约：

- 场景卡片加载：前端 init() 拉取 /api/scene-cards → 左侧导航渲染场景/子场景/阶段
- 对话收集（规则降级）：/api/chat 在 LLM 不可用时返回 llm:false，前端 _fallbackAsk
  规则式把用户输入填入首个空字段
- 卡片收集（chip）：select/multi 字段用 .chip 点选 → 字段填充 + 进度条 + 汇总更新
- 生成预览：generateResume() 纯客户端聚合 → 预览模态框 + 结构化小节
- 导出 docx：downloadWord() POST /api/export → 真实 docx 字节流
- 导入 docx：uploadResume() 上传 .docx → /api/import-docx 解析 → 还原排版样式
- 导入 xlsx：上传 .xlsx → /api/import-xlsx（openpyxl 单元格样式 → HTML 表格）
- 导入 pdf：上传 .pdf → /api/import-pdf（pdfplumber 字符级 → CSS/HTML）
- 导入 image：上传 .png → /api/import-image（OCR mock，验证前端分派契约）

后端复用 create_app()（真实 scene_cards.json + 真实 python-docx），LLM 用 monkeypatch
替换为确定性降级，因此本用例是真实浏览器渲染 + 真实 HTTP，秒级完成、无外部依赖。

运行：`python -m pytest product/tests/test_e2e_browser.py -v`
（未安装 playwright 或 chromium 时自动跳过，不影响其余 439 用例。）
"""

from __future__ import annotations

import glob
import os
import sys
import threading

import pytest

_HERE = os.path.dirname(os.path.abspath(__file__))
_PRODUCT = os.path.dirname(_HERE)  # product 目录
_SRC = os.path.join(_PRODUCT, "src")
if _SRC not in sys.path:
    sys.path.insert(0, _SRC)

from werkzeug.serving import make_server  # noqa: E402

from resume_product.api import create_app  # noqa: E402
from resume_product import llm_client  # noqa: E402

playwright_sync = pytest.importorskip(
    "playwright.sync_api", reason="需要 playwright（pip install playwright）")
from playwright.sync_api import expect, sync_playwright  # noqa: E402


def _browser_executable() -> str | None:
    """定位可用的 Chromium/Chrome/Edge（自装 chromium 优先，系统浏览器兜底）。"""
    local = os.environ.get("LOCALAPPDATA", "")
    if local:
        for pat in (
            os.path.join(local, "ms-playwright", "chromium-*", "chrome-win64", "chrome.exe"),
            os.path.join(local, "ms-playwright", "chromium-*", "chrome-win", "chrome.exe"),
        ):
            for exe in glob.glob(pat):
                if os.path.exists(exe):
                    return exe
    for pth in (
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    ):
        if os.path.exists(pth):
            return pth
    return None


pytestmark = pytest.mark.skipif(
    _browser_executable() is None,
    reason="无可用浏览器（playwright install chromium 或安装 Chrome/Edge）")


@pytest.fixture(scope="module")
def live_url():
    """模块级：后台线程启动一次真实 Flask HTTP 服务（临时端口）。"""
    app = create_app()
    server = make_server("127.0.0.1", 0, app, threaded=True)
    port = server.server_port
    t = threading.Thread(target=server.serve_forever, daemon=True)
    t.start()
    yield f"http://127.0.0.1:{port}"
    server.shutdown()
    t.join(timeout=5)


@pytest.fixture(scope="module")
def browser():
    p = sync_playwright().start()
    exe = _browser_executable()
    b = p.chromium.launch(headless=True, executable_path=exe)
    try:
        yield b
    finally:
        b.close()
        p.stop()


@pytest.fixture()
def page(browser, live_url):
    pg = browser.new_page()
    pg.set_default_timeout(10000)
    yield pg
    pg.close()


# ── E2E-1：场景卡片加载（/api/scene-cards → 左侧导航渲染）───────────
def test_scene_cards_load(page, live_url):
    page.goto(live_url + "/")
    # 场景头（每个场景一个 .scene-head）
    page.wait_for_selector(".scene-head")
    heads = page.locator(".scene-head")
    expect(heads.first).to_be_visible()
    assert heads.count() >= 4, "场景卡片应含多个场景（保研/考研/出国/求职等）"
    # 至少能看到「求职」与「保研升学」两个场景
    nav_text = page.locator("#nav").inner_text()
    assert "求职" in nav_text
    assert "保研升学" in nav_text
    # 阶段节点已渲染（.stage）
    expect(page.locator(".nav .stage").first).to_be_visible()


# ── E2E-2：对话收集（LLM 降级 → 规则式填入首个空字段）──────────────
def test_dialog_send_degrades_to_rule(page, live_url, monkeypatch):
    # LLM 不可用（返回空串）→ /api/chat 返回 llm:false → 前端走 _fallbackAsk
    monkeypatch.setattr(llm_client, "chat", lambda *a, **k: "")
    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")

    # 选「求职 / 技术岗·后端开发 / 基本信息」阶段（首个字段 name=姓名，text）
    page.click('.nav .stage[data-scene="job"][data-sub="backend"][data-stage="basic"]')
    page.wait_for_selector(".card-body .field")
    # 首个字段应为「姓名」（required text）
    first_label = page.locator(".card-body .field .f-label").first.inner_text()
    assert "姓名" in first_label

    # 发送一句话 → 规则式填入「姓名」
    page.fill("#input", "张三")
    page.click(".composer .btn")  # 发送
    page.wait_for_function(
        "document.querySelector('.card-body .field .f-value') && "
        "document.querySelector('.card-body .field').classList.contains('filled')")
    expect(page.locator(".card-body .field").first).to_have_class(
        __import__("re").compile(r"\bfilled\b"))
    expect(page.locator(".card-body .field .f-value").first).to_contain_text("张三")
    # 进度：1/N 已收集
    expect(page.locator("#cardProgress")).to_contain_text("已收集")


# ── E2E-3：卡片 chip 收集（multi 字段点选 → 填充/进度/汇总更新）────
def test_chip_collection_multi_field(page, live_url):
    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")

    # 「求职 / 后端 / 核心技能」阶段（首个字段 languages=语言，multi）
    page.click('.nav .stage[data-scene="job"][data-sub="backend"][data-stage="core"]')
    page.wait_for_selector(".card-body .field")
    expect(page.locator(".card-body .field")).to_have_count(4)  # 语言/框架/数据库/中间件

    # 首个字段为「语言」，askNext 渲染 chip 选项
    first_label = page.locator(".card-body .field .f-label").first.inner_text()
    assert "语言" in first_label
    expect(page.locator(".msg .chip").first).to_be_visible()

    # 点选 chip「Java」→ 字段填充 + 进度 1/4
    page.locator(".msg .chip", has_text="Java").first.click()
    expect(page.locator(".card-body .field").first).to_have_class(
        __import__("re").compile(r"\bfilled\b"))
    expect(page.locator(".card-body .field .f-value").first).to_contain_text("Java")
    expect(page.locator("#cardProgress")).to_have_text("1/4 已收集")

    # 再点选「Go」→ multi 累积「Java、Go」（同一字段，进度仍 1/4）
    page.locator(".msg .chip", has_text="Go").first.click()
    expect(page.locator(".card-body .field .f-value").first).to_contain_text("Java、Go")
    expect(page.locator("#cardProgress")).to_have_text("1/4 已收集")

    # 导航圆点进入 partial 态（部分收集）
    expect(page.locator("#dot-job-backend-core")).to_have_class(
        __import__("re").compile(r"\bpartial\b"))


# ── E2E-4：生成预览（客户端聚合 → 模态框 + 结构化小节）─────────────
def test_generate_preview(page, live_url):
    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")
    page.click('.nav .stage[data-scene="job"][data-sub="backend"][data-stage="core"]')
    page.wait_for_selector(".card-body .field")
    # 先填一点数据（语言=Java）
    page.locator(".msg .chip", has_text="Java").first.click()
    expect(page.locator(".card-body .field .f-value").first).to_contain_text("Java")

    # 点「生成简历」→ 预览模态框打开
    page.click("#btnGen")
    expect(page.locator("#previewModal")).to_have_class(
        __import__("re").compile(r"\bopen\b"))
    body = page.locator("#previewBody")
    expect(body).to_be_visible()
    assert "核心技能" in body.inner_text()
    assert "语言：Java" in body.inner_text()

    # 切换模板：现代/极简（纯客户端 class 切换，不崩）
    page.click('.tpl-btn[data-tpl="modern"]')
    expect(page.locator("#previewBody")).to_have_class(
        __import__("re").compile(r"\bpv-modern\b"))
    page.click('.tpl-btn[data-tpl="minimal"]')
    expect(page.locator("#previewBody")).to_have_class(
        __import__("re").compile(r"\bpv-minimal\b"))

    # 关闭预览
    page.click(".modal-head .close")
    expect(page.locator("#previewModal")).not_to_have_class(
        __import__("re").compile(r"\bopen\b"))


# ── E2E-5：导出 docx（/api/export 真实字节流契约）───────────────────
def test_export_docx_contract(page, live_url):
    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")
    page.click('.nav .stage[data-scene="job"][data-sub="backend"][data-stage="core"]')
    page.wait_for_selector(".card-body .field")
    page.locator(".msg .chip", has_text="Java").first.click()
    expect(page.locator(".card-body .field .f-value").first).to_contain_text("Java")

    # generateResume() 设置 window._resumeText（导出数据源）
    page.click("#btnGen")
    page.wait_for_function("typeof window._resumeText === 'string' && window._resumeText.length > 0")

    # 直接走 downloadWord 的真实 fetch 契约（返回 docx 字节流）
    result = page.evaluate("""async () => {
      const r = await fetch('/api/export', {
        method: 'POST', headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({experiences: [], target_role: '后端工程师',
          resume_text: window._resumeText || '', format: 'docx'})});
      const buf = await r.arrayBuffer();
      return {status: r.status, ok: r.ok, len: buf.byteLength,
              ct: r.headers.get('content-type') || ''};
    }""")
    assert result["ok"] is True, result
    assert result["status"] == 200
    assert result["len"] > 500, "docx 应为非空 zip 字节流"
    # content-type 为 docx（application/vnd.openxmlformats-...）
    assert "wordprocessingml" in result["ct"] or "octet-stream" in result["ct"], result["ct"]


# ── E2E-6：导入 docx（上传 → /api/import-docx 解析 → 还原排版）──────
def test_import_docx_via_upload(page, live_url, tmp_path):
    import docx
    d = docx.Document()
    d.add_paragraph("张三")
    d.add_paragraph("求职意向：后端开发工程师")
    src = tmp_path / "resume.docx"
    d.save(str(src))

    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")

    # 上传文件（前端 uploadResume → /api/import-docx）
    page.set_input_files('input[type="file"]', str(src))
    page.wait_for_function(
        "document.getElementById('uploadMsg').textContent.includes('已还原排版')")
    expect(page.locator("#uploadMsg")).to_contain_text("已还原排版")
    # 对话区出现「已识别你的简历」提示（含段落数）
    expect(page.locator(".msg.ai").last).to_contain_text("已识别你的简历")


# ── E2E-7：导入 xlsx（上传 → /api/import-xlsx → 还原表格排版）─────────
def test_import_xlsx_via_upload(page, live_url, tmp_path):
    """Excel 简历上传走真实 /api/import-xlsx 契约（openpyxl 单元格样式 → HTML 表格）。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "王小明 简历"
    ws.merge_cells("A1:C1")
    ws["A1"].font = Font(name="黑体", size=16, bold=True, color="FF1F3A5F")
    ws["A1"].fill = PatternFill("solid", fgColor="FFE8EDFF")
    ws["A2"], ws["B2"], ws["C2"] = "技能类别", "技能", "水平"
    ws["A3"], ws["B3"], ws["C3"] = "产品", "需求分析", "精通"
    src = tmp_path / "resume.xlsx"
    wb.save(str(src))

    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")

    # 上传 xlsx（前端 uploadResume 按扩展名路由到 /api/import-xlsx）
    page.set_input_files('input[type="file"]', str(src))
    page.wait_for_function(
        "document.getElementById('uploadMsg').textContent.includes('已还原排版')")
    expect(page.locator("#uploadMsg")).to_contain_text("已还原排版")
    # 对话区提示含行数（xlsx 的 meta.paragraphs = 行数）
    expect(page.locator(".msg.ai").last).to_contain_text("已识别你的简历")


# ── E2E-8：导入 pdf（上传 → /api/import-pdf → pdfplumber 还原排版）─────
def test_import_pdf_via_upload(page, live_url, tmp_path):
    """PDF 简历上传走真实 /api/import-pdf 契约（pdfplumber 字符级 → CSS/HTML）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor
    src = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(src), pagesize=A4)
    W, H = A4
    c.setFillColor(HexColor("#1F3A5F"))
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(W / 2, H - 80, "Resume")
    c.setFillColor(HexColor("#000000"))
    c.setFont("Helvetica", 11)
    c.drawString(50, H - 140, "Work Experience")
    c.drawString(50, H - 170, "Engineer at Google")
    c.drawString(350, H - 140, "Skills")
    c.drawString(350, H - 170, "Python, SQL")
    c.save()

    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")

    # 上传 pdf（前端 uploadResume 按扩展名路由到 /api/import-pdf）
    page.set_input_files('input[type="file"]', str(src))
    page.wait_for_function(
        "document.getElementById('uploadMsg').textContent.includes('已还原排版')")
    expect(page.locator("#uploadMsg")).to_contain_text("已还原排版")
    expect(page.locator(".msg.ai").last).to_contain_text("已识别你的简历")


# ── E2E-9：导入 image（mock OCR → 验证前端 /api/import-image 契约）────
def test_import_image_via_upload_mocked(page, live_url, tmp_path, monkeypatch):
    """图片简历上传走真实 /api/import-image 契约（OCR 用 mock 跳过 rapidocr 慢路径）。"""
    from resume_product.render import image_import
    monkeypatch.setattr(image_import, "import_resume_image", lambda _path: {
        "html": '<div style="font-size:14pt;text-align:center;">王小明</div>',
        "css": "@page { size: A4; margin: 18mm 16mm; }",
        "meta": {"paragraphs": 2, "columns": 1, "css_length": 42},
        "filename": "resume.png",
    })
    # 最小 PNG（内容不参与：OCR 已被 mock）
    src = tmp_path / "resume.png"
    src.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)

    page.goto(live_url + "/")
    page.wait_for_selector(".scene-head")

    # 上传 png（前端 uploadResume 按扩展名路由到 /api/import-image）
    page.set_input_files('input[type="file"]', str(src))
    page.wait_for_function(
        "document.getElementById('uploadMsg').textContent.includes('已还原排版')")
    expect(page.locator("#uploadMsg")).to_contain_text("已还原排版")
    expect(page.locator(".msg.ai").last).to_contain_text("已识别你的简历")
