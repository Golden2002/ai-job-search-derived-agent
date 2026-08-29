# -*- coding: utf-8 -*-
"""resume_product.render.docx_import —— Word 简历排版解析 → 定制 CSS（完美还原排版）。

用户核心诉求：上传自己的简历（一般为 Word 文档），工具根据其排版自动生成定制 CSS，
使渲染结果**完美、精准还原原文档的排版**——包括段落（字体/字号/对齐/缩进/间距/颜色）、
**表格、分栏、图片（图形元素）**等复杂排版元素。

对外 API：
- parse_docx(docx_path)             → 结构化块（段落/表格/图片，按文档顺序）+ 分栏数
- import_docx_resume(docx_path)     → {html, css, meta} 一站式（css 可直接作 custom_css）

设计：
- 按 body 元素顺序遍历（paragraph / table 交错），还原文档真实顺序
- 段落：run 字体 + 段落格式 → 内联样式（跨 run 字体差异用 span 保留）
- 表格：行/列/单元格文本 → HTML <table> + 单元格样式
- 图片：inline shape 提取 → base64 data URI <img>（保留图形元素）
- 分栏：section 的 w:cols → CSS column-count（双栏简历还原）

依赖：python-docx（resume_extract 可选依赖，已在 pyproject.toml 声明）。
"""
from __future__ import annotations

import base64
from typing import Any, Dict, List, Tuple

try:  # 可选依赖：python-docx（resume_extract）
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    qn = None  # type: ignore


def _pt(v) -> float:
    """Length 对象 → pt 浮点（None → 0.0）。"""
    try:
        return float(v.pt) if v is not None else 0.0
    except Exception:
        return 0.0


def _rgb_to_hex(rgb) -> str:
    """RGBColor → '#RRGGBB'（None → ''）。"""
    try:
        return f"#{rgb}" if rgb is not None else ""
    except Exception:
        return ""


def _align_css(a) -> str:
    if a is None:
        return ""
    s = str(a)
    if "CENTER" in s:
        return "center"
    if "RIGHT" in s:
        return "right"
    if "JUSTIFY" in s:
        return "justify"
    return "left"


def _line_spacing_css(ls) -> str:
    if ls is None:
        return ""
    try:
        if isinstance(ls, (int, float)):
            return f"line-height:{ls}"
        return f"line-height:{ls.pt:.1f}pt"
    except Exception:
        return ""


def _esc(s) -> str:
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ─────────────────────────────────────
# run / 段落解析（保留原逻辑）
# ─────────────────────────────────────
def _run_to_dict(r) -> Dict[str, Any]:
    f = r.font
    return {
        "text": r.text,
        "font_name": f.name,
        "size_pt": round(f.size.pt, 1) if f.size is not None else None,
        "bold": bool(f.bold),
        "italic": bool(f.italic),
        "underline": bool(f.underline),
        "color": _rgb_to_hex(f.color.rgb if f.color else None),
    }


def _paragraph_images(document, p_el) -> List[str]:
    """段落内嵌图片 → base64 data URI 列表（图形元素保留）。"""
    images = []
    try:
        for blip in p_el.findall(".//" + qn("a:blip")):
            rid = blip.get(qn("r:embed"))
            if not rid:
                continue
            part = document.part.related_parts.get(rid)
            if part is None:
                continue
            data = base64.b64encode(part.blob).decode("ascii")
            ext = (part.partname.ext or ".png").lstrip(".")
            if ext.lower() in ("jpg", "jpeg", "png", "gif", "bmp", "webp"):
                mime = "jpeg" if ext.lower() in ("jpg",) else ext.lower()
                images.append(f"data:image/{mime};base64,{data}")
    except Exception:
        pass
    return images


def _parse_paragraph(document, p_el) -> Dict[str, Any]:
    """w:p 元素 → 段落 dict（含 runs + 排版 + 图片 + 列表类型）。"""
    from docx.text.paragraph import Paragraph
    p = Paragraph(p_el, document)
    text = p.text.strip()
    if not text and not _paragraph_images(document, p_el):
        return None  # 空段落且无图 → 跳过
    fmt = p.paragraph_format
    # 列表类型：pStyle 名称（ListBullet/ListNumber）→ bullet/decimal
    style_name = (p.style.name or "") if p.style else ""
    list_type = None
    if "bullet" in style_name.lower():
        list_type = "bullet"
    elif "number" in style_name.lower() or "decimal" in style_name.lower():
        list_type = "decimal"
    return {
        "type": "paragraph",
        "text": text,
        "runs": [_run_to_dict(r) for r in p.runs],
        "images": _paragraph_images(document, p_el),
        "alignment": _align_css(p.alignment),
        "first_line_indent_pt": round(_pt(fmt.first_line_indent), 1),
        "left_indent_pt": round(_pt(fmt.left_indent), 1),
        "space_before_pt": round(_pt(fmt.space_before), 1),
        "space_after_pt": round(_pt(fmt.space_after), 1),
        "line_spacing": _line_spacing_css(fmt.line_spacing),
        "list_type": list_type,
    }


def _cell_text(tc) -> str:
    """w:tc → 单元格文本（拼接所有段落，换行分隔）。"""
    parts = []
    for p in tc.findall(qn("w:p")):
        t = "".join(node.text or "" for node in p.iter(qn("w:t")))
        if t.strip():
            parts.append(t)
    return "\n".join(parts)


def _parse_table(tbl) -> Dict[str, Any]:
    """w:tbl 元素 → 表格 dict（行/列/单元格文本 + 合并单元格 gridSpan/vMerge）。"""
    rows = []
    for tr in tbl.findall(qn("w:tr")):
        cells = []
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            grid_span = 1
            v_merge = None
            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    try:
                        grid_span = int(gs.get(qn("w:val"), "1"))
                    except Exception:
                        grid_span = 1
                vm = tcPr.find(qn("w:vMerge"))
                if vm is not None:
                    # restart = 合并起始，缺省 val = 继续合并
                    v_merge = vm.get(qn("w:val")) or "continue"
            cells.append({"text": _cell_text(tc), "grid_span": grid_span,
                          "v_merge": v_merge})
        rows.append(cells)
    ncols = max((sum(c.get("grid_span", 1) for c in r) for r in rows), default=0)
    return {"type": "table", "rows": rows, "ncols": ncols}


def _section_columns(document) -> int:
    """分栏数（section w:cols 的 num 属性，默认 1）。"""
    try:
        if document.sections:
            sect = document.sections[0]._sectPr
            cols = sect.find(qn("w:cols"))
            if cols is not None:
                n = cols.get(qn("w:num"))
                if n:
                    return int(n)
    except Exception:
        pass
    return 1


def _parse_headers_footers(document) -> Dict[str, List[str]]:
    """提取页眉/页脚文本（简历常见：页眉放姓名/联系方式）。"""
    headers: List[str] = []
    footers: List[str] = []
    for section in document.sections:
        try:
            h = section.header
            if h is not None and not h.is_linked_to_previous:
                for p in h.paragraphs:
                    if p.text.strip():
                        headers.append(p.text.strip())
        except Exception:
            pass
        try:
            f = section.footer
            if f is not None and not f.is_linked_to_previous:
                for p in f.paragraphs:
                    if p.text.strip():
                        footers.append(p.text.strip())
        except Exception:
            pass
    return {"headers": headers, "footers": footers}


# ─────────────────────────────────────
# 解析入口
# ─────────────────────────────────────
def parse_docx(docx_path: str) -> Dict[str, Any]:
    """解析 Word 文档 → 结构化块（段落/表格，按文档顺序）+ 分栏数。"""
    import docx

    document = docx.Document(docx_path)
    body = document.element.body
    blocks: List[Dict[str, Any]] = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = _parse_paragraph(document, child)
            if p is not None:
                blocks.append(p)
        elif child.tag == qn("w:tbl"):
            blocks.append(_parse_table(child))
    return {
        "blocks": blocks,
        "columns": _section_columns(document),
        "paragraphs": [b for b in blocks if b.get("type") == "paragraph"],
        "tables": [b for b in blocks if b.get("type") == "table"],
        "headers": _parse_headers_footers(document)["headers"],
        "footers": _parse_headers_footers(document)["footers"],
    }


# ─────────────────────────────────────
# 样式生成
# ─────────────────────────────────────
def _paragraph_style(p: Dict[str, Any]) -> str:
    decls = []
    runs = p.get("runs") or []
    if runs:
        r0 = runs[0]
        if r0.get("font_name"):
            decls.append(f"font-family:'{r0['font_name']}',serif")
        if r0.get("size_pt"):
            decls.append(f"font-size:{r0['size_pt']:.1f}pt")
    if p.get("alignment"):
        decls.append(f"text-align:{p['alignment']}")
    if p.get("first_line_indent_pt"):
        decls.append(f"text-indent:{p['first_line_indent_pt']:.1f}pt")
    if p.get("left_indent_pt"):
        decls.append(f"margin-left:{p['left_indent_pt']:.1f}pt")
    if p.get("space_before_pt"):
        decls.append(f"margin-top:{p['space_before_pt']:.1f}pt")
    if p.get("space_after_pt"):
        decls.append(f"margin-bottom:{p['space_after_pt']:.1f}pt")
    if p.get("line_spacing"):
        decls.append(p["line_spacing"])
    return ";".join(decls)


def _run_html(r: Dict[str, Any]) -> str:
    style = []
    if r.get("font_name"):
        style.append(f"font-family:'{r['font_name']}',serif")
    if r.get("size_pt"):
        style.append(f"font-size:{r['size_pt']:.1f}pt")
    if r.get("bold"):
        style.append("font-weight:bold")
    if r.get("italic"):
        style.append("font-style:italic")
    if r.get("underline"):
        style.append("text-decoration:underline")
    if r.get("color"):
        style.append(f"color:{r['color']}")
    text = _esc(r.get("text", ""))
    if style:
        return f'<span style="{";".join(style)}">{text}</span>'
    return text


def _signature(p: Dict[str, Any]) -> Tuple:
    runs = tuple(
        (r.get("font_name"), r.get("size_pt"), r.get("bold"),
         r.get("italic"), r.get("underline"))
        for r in p.get("runs", [])
    )
    return (
        runs,
        p.get("alignment", ""),
        p.get("first_line_indent_pt", 0),
        p.get("left_indent_pt", 0),
        p.get("space_before_pt", 0),
        p.get("space_after_pt", 0),
        p.get("line_spacing", ""),
    )


# ─────────────────────────────────────
# HTML / CSS 生成
# ─────────────────────────────────────
def _block_html(b: Dict[str, Any], list_prefix: str = "") -> str:
    """单个块 → HTML（list_prefix 为列表符号/编号前缀）。"""
    if b.get("type") == "table":
        rows_html = []
        for row in b.get("rows", []):
            cells_html = ""
            for c in row:
                if isinstance(c, dict):
                    text = c.get("text", "")
                    gs = c.get("grid_span", 1)
                    vm = c.get("v_merge")
                else:  # 兼容旧格式（纯字符串）
                    text = c
                    gs, vm = 1, None
                if vm == "continue":
                    continue  # 垂直合并的后续单元格跳过
                colspan = f' colspan="{gs}"' if gs > 1 else ""
                cells_html += (f'<td style="border:1px solid #999;padding:4pt 6pt"'
                               f'{colspan}>{_esc(text)}</td>')
            rows_html.append(f"<tr>{cells_html}</tr>")
        return ('<table style="border-collapse:collapse;width:100%;'
                'margin:6pt 0;">' + "".join(rows_html) + "</table>")
    # 段落
    runs_html = "".join(_run_html(r) for r in b.get("runs", []))
    imgs_html = "".join(
        f'<img src="{img}" style="max-width:100%;height:auto;" alt="">'
        for img in b.get("images", []))
    content = runs_html or _esc(b.get("text", ""))
    prefix_html = (f'<span style="margin-right:5pt">{_esc(list_prefix)}</span>'
                   if list_prefix else "")
    return f'<div style="{_paragraph_style(b)}">{imgs_html}{prefix_html}{content}</div>'


def import_docx_resume(docx_path: str) -> Dict[str, Any]:
    """上传的 Word 简历 → {html, css, meta}（CSS 完美还原原排版，含表格/分栏/图片）。"""
    parsed = parse_docx(docx_path)
    blocks = parsed.get("blocks", [])
    columns = parsed.get("columns", 1)

    # HTML：按文档顺序渲染块（段落/表格/图片 + 列表符号/编号还原）
    html_parts = []
    decimal_counter = 0
    for b in blocks:
        list_type = b.get("list_type") if b.get("type") == "paragraph" else None
        if list_type == "decimal":
            decimal_counter += 1
            html_parts.append(_block_html(b, f"{decimal_counter}. "))
        else:
            decimal_counter = 0  # 非 decimal（bullet/正文/表格）打断编号连续性
            prefix = "• " if list_type == "bullet" else ""
            html_parts.append(_block_html(b, prefix))
    html = "\n".join(html_parts)

    # 页眉/页脚（简历常见：页眉放姓名/联系方式）
    headers = parsed.get("headers", [])
    footers = parsed.get("footers", [])
    header_html = ""
    footer_html = ""
    if headers:
        header_html = ('<div style="text-align:center;font-size:9pt;color:#555;'
                       'border-bottom:1px solid #ccc;padding-bottom:4pt;'
                       'margin-bottom:8pt">'
                       + " | ".join(_esc(h) for h in headers) + "</div>")
    if footers:
        footer_html = ('<div style="text-align:center;font-size:9pt;color:#555;'
                       'border-top:1px solid #ccc;padding-top:4pt;'
                       'margin-top:8pt">'
                       + " | ".join(_esc(f) for f in footers) + "</div>")
    html = "\n".join(x for x in [header_html, html, footer_html] if x)

    # CSS：去重段落类 + 分栏
    css_parts = [
        "@page { size: A4; margin: 22mm 22mm 20mm 22mm; }",
        "body { font-family: 'SimSun','Songti SC','Noto Serif CJK SC',serif; }",
    ]
    if columns > 1:
        css_parts.append(
            f".docx-columns {{ column-count:{columns}; column-gap:8mm; }}")
    sig_to_idx: Dict[Tuple, int] = {}
    for p in parsed.get("paragraphs", []):
        sig = _signature(p)
        if sig not in sig_to_idx:
            sig_to_idx[sig] = len(sig_to_idx)
            css_parts.append(
                f".docx-p-{sig_to_idx[sig]} {{ {_paragraph_style(p)}; }}")
    css = "\n".join(css_parts)

    fonts = sorted({
        r.get("font_name", "")
        for p in parsed.get("paragraphs", [])
        for r in p.get("runs", [])
        if r.get("font_name")
    })
    meta = {
        "paragraphs": len(parsed.get("paragraphs", [])),
        "tables": len(parsed.get("tables", [])),
        "images": sum(len(p.get("images", [])) for p in parsed.get("paragraphs", [])),
        "columns": columns,
        "headers": len(parsed.get("headers", [])),
        "footers": len(parsed.get("footers", [])),
        "distinct_styles": len(sig_to_idx),
        "css_length": len(css),
        "fonts": fonts,
    }
    return {"html": html, "css": css, "meta": meta, "parsed": parsed}
