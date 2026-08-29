# -*- coding: utf-8 -*-
"""resume_product.core — 通用简历引擎（W4-W5 ⭐）。

复用 medical-resume-agent 引擎（claim_gate/confirmation_gate/experience_draft/
bullet_composer）+ 通用化改造（去医学化 + 全行业 Role Pack）。

架构：
- 经历采集：LLM 从原始文本提取事实卡（主张必须引用原文——claim gate 思想）
- 定向表达：Role Pack 适配（能力重排/动词优化/句式模板）
- 简历生成：结构化事实卡 → markdown/html（可扩展 LaTeX/PDF）
- 可注入 chat_fn（PAEG 生态——外部智能体接入）
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from . import llm_client

# 语言规范模块（PAEG 工具生态 14.1——独立仓库同步接入）
try:
    from paeg_lang_style import gate_short as _lang_gate
    from paeg_lang_style import fix_known_gaffes as _lang_fix
    _HAS_LANG_STYLE = True
except ImportError:
    _HAS_LANG_STYLE = False

    def _lang_gate(text: str, context: str = "") -> str:
        return text

    def _lang_fix(text: str) -> str:
        return text


def _lang_l0(text: str) -> str:
    """L0 语言校对：gate_short（快路径） + fix_known_gaffes 兜底收口。

    缺失 paeg_lang_style 时优雅降级为原文（不抛异常、不阻塞生成链路）。
    """
    if not text or not _HAS_LANG_STYLE:
        return text
    try:
        out = _lang_gate(text)
        return _lang_fix(out if isinstance(out, str) else text)
    except Exception:
        return _lang_fix(text)

# §3.116 ⭐ R-09 融合修复：medical-resume-agent 引擎已**本地移植**（不重复造轮子）——
# claim_gate/confirmation_gate/canonical_experience/multi_version/capability_taxonomy
# 均已落地为 product/src/resume_product/ 下同名模块（见 claim_gate.py 等）。
# 原 _MEDICAL_SRC 外部路径引用（指向不存在的 "medical-resume-agent" 目录）为死代码，
# 已移除——本地移植版为权威实现，不依赖 14.6 运行时路径。

# 数据目录（product/data/——role_packs + capability_tags）
_DATA_DIR = Path(__file__).resolve().parent.parent.parent / "data"


def _load_json(name: str) -> Dict[str, Any]:
    p = _DATA_DIR / name
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _load_role_packs() -> Dict[str, Dict[str, Any]]:
    """加载通用 Role Pack（data/role_packs/*.json）。"""
    packs: Dict[str, Dict[str, Any]] = {}
    d = _DATA_DIR / "role_packs"
    if d.is_dir():
        for f in d.glob("*.json"):
            try:
                packs[f.stem] = json.loads(f.read_text(encoding="utf-8"))
            except Exception:
                continue
    return packs


def _default_chat(sys_p: str, usr_p: str) -> str:
    """默认 chat_fn：DeepSeek LLM（失败返回空串 → 调用方降级确定性模式）。"""
    return llm_client.chat(sys_p, usr_p)


class ResumeEngine:
    """通用简历引擎（可注入 LLM——PAEG 生态 ⭐）。"""

    def __init__(self, chat_fn: Optional[Callable] = None,
                 role_pack: str = "tech_v1"):
        self.chat_fn = chat_fn or _default_chat
        self.role_pack = role_pack

    def enrich(self, raw_text: str) -> List[Dict[str, str]]:
        """经历文本 → 结构化事实卡（主张校验思想）。

        总是尝试 LLM（默认 chat_fn 即 DeepSeek，失败/空串时降级）。
        LLM 拆解：口语规范化 + 结构化事实卡（引用原文 + 保留量化数据）。
        """
        if not raw_text or not raw_text.strip():
            return []
        # LLM 提取（主张校验——要求引用原文）
        sys_p = ("你是经历拆解助手。从用户提供的经历描述中提取结构化事实卡。"
                 "用户输入往往含口语化、自嘲式表述（如『瞎调参数』『充数』『spice monkey』），"
                 "你必须先做口语规范化：把口语/自嘲表述转为书面、客观、规范的简历语言，"
                 "同时保留真实信息与量化数据，绝不编造、绝不夸大。"
                 "每张事实卡必须：1) claim 为规范化后的书面表述 2) 标注可验证的证据 "
                 "3) 保留量化数据 4) quote 引用原文关键片段。"
                 '输出 JSON 数组：[{"claim": "规范化主张", "evidence": "证据", "quote": "原文引用"}]')
        raw = self.chat_fn(sys_p, f"经历描述：{raw_text}")
        if raw and raw.strip():
            try:
                import re
                m = re.search(r"\[.*\]", raw or "", re.S)
                if m:
                    data = json.loads(m.group(0))
                    facts = [d for d in data if isinstance(d, dict) and d.get("claim")]
                    if facts:
                        return facts
            except Exception:
                pass
        return self._heuristic_extract(raw_text)

    def _heuristic_extract(self, raw_text: str) -> List[Dict[str, str]]:
        """确定性提取（无 LLM 兜底）：按句拆事实卡。"""
        import re
        sentences = re.split(r"[。！？!?；;]", raw_text)
        facts = []
        for s in sentences:
            s = s.strip()
            if len(s) < 4:
                continue
            facts.append({"claim": s, "evidence": s,
                          "quote": s, "source": "heuristic"})
        return facts

    def compose(self, facts: List[Dict[str, str]],
                target_role: str = "") -> str:
        """定向表达：Role Pack 适配 → markdown 简历。"""
        pack = _load_role_packs().get(self.role_pack, {})
        label = pack.get("label", "通用")
        lines = [f"# {target_role or '个人简历'}", ""]
        if pack.get("priorities"):
            lines.append(f"**适配方向**：{label}（{', '.join(pack['priorities'])}）")
            lines.append("")
        for i, fact in enumerate(facts[:10], 1):
            claim = fact.get("claim", "")
            ev = fact.get("evidence", "")
            if claim:
                # 语言规范模块（14.1 paeg_lang_style）：gate_short/fix_known_gaffes L0 校对后再输出
                claim = _lang_l0(claim)
                lines.append(f"{i}. {claim}")
                if ev and ev != claim:
                    lines.append(f"   - 证据：{_lang_l0(ev)}")
        return "\n".join(lines)

    def to_html(self, facts: List[Dict[str, str]],
                target_role: str = "") -> str:
        """结构化 → HTML 简历。"""
        md = self.compose(facts, target_role)
        import html as _h
        title = _h.escape(target_role or "个人简历")
        body = ""
        for line in md.split("\n"):
            if line.startswith("# "):
                body += f"<h1>{_h.escape(line[2:])}</h1>"
            elif line.startswith("**"):
                # 去掉加粗标记（**适配方向**：…），避免 ** 泄入 HTML
                body += f"<p><strong>{_h.escape(line.replace('**', ''))}</strong></p>"
            elif line.strip().startswith("- "):
                body += f"<li>{_h.escape(line.strip()[2:])}</li>"
            elif line.strip() and line.strip()[0].isdigit() and "." in line[:4]:
                body += f"<p>{_h.escape(line)}</p>"
            elif line.strip():
                body += f"<p>{_h.escape(line)}</p>"
        return f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">
<title>{title}</title>
<style>body{{font-family:'SimSun',serif;max-width:800px;margin:40px auto;line-height:1.7}}
h1{{text-align:center;border-bottom:1px solid #333;padding-bottom:8px}}</style>
</head><body>{body}</body></html>"""

    def to_docx(self, facts: List[Dict[str, str]],
                target_role: str = "", out_path: str = "") -> str:
        """结构化 → Word 简历（python-docx，专业排版）。

        Returns: 生成的 .docx 文件路径。
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # python-docx 未装 → 降级生成 .doc（纯文本）
            import tempfile
            md = self.compose(facts, target_role)
            if not out_path:
                out_path = os.path.join(tempfile.gettempdir(),
                                        "resume_plain.txt")
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(md)
            return out_path

        import tempfile
        doc = Document()
        # 页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 标题（居中）
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(target_role or "个人简历")
        run.font.size = Pt(20)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)

        # 适配方向
        pack = _load_role_packs().get(self.role_pack, {})
        label = pack.get("label", "通用")
        if pack.get("priorities"):
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sub.add_run(f"适配方向：{label}（{', '.join(pack['priorities'])}）")
            sr.font.size = Pt(10)
            sr.font.color.rgb = RGBColor(0x5C, 0x53, 0x4A)

        # 分隔线
        doc.add_paragraph()

        # 经历（事实卡）
        for i, fact in enumerate(facts[:10], 1):
            claim = fact.get("claim", "")
            ev = fact.get("evidence", "")
            if not claim:
                continue
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {claim}")
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)
            if ev and ev != claim:
                ep = doc.add_paragraph()
                ep.paragraph_format.left_indent = Cm(0.6)
                er = ep.add_run(f"证据：{ev}")
                er.font.size = Pt(10)
                er.font.color.rgb = RGBColor(0x5C, 0x53, 0x4A)
            p.paragraph_format.space_after = Pt(4)

        if not out_path:
            out_path = os.path.join(tempfile.gettempdir(), "resume.docx")
        doc.save(out_path)
        return out_path

    def to_pdf(self, facts: List[Dict[str, str]],
               target_role: str = "", out_path: str = "") -> str:
        """结构化 → PDF（固定资产渲染：render/resume.css + render_pdf.py）。

        Returns: 生成的 .pdf 文件路径。
        """
        import tempfile
        # 1. 生成 markdown
        md = self.compose(facts, target_role)
        # 2. 固定资产渲染（render/resume.css + render_pdf.py——Playwright + Chrome）
        from .render import render_markdown_to_pdf
        if not out_path:
            out_path = os.path.join(tempfile.gettempdir(), "resume.pdf")
        try:
            return render_markdown_to_pdf(md, out_path)
        except Exception:
            # 降级：Playwright 不可用 → 返回 markdown（标记）
            with open(out_path.replace(".pdf", ".md"), "w", encoding="utf-8") as f:
                f.write(md)
            return out_path.replace(".pdf", ".md")


# ═══════════════════════════════════════════════════════════
# 便捷 API（MCP 工具对应）
# ═══════════════════════════════════════════════════════════
_engine = ResumeEngine()


def enrich_experience(raw_text: str) -> List[Dict[str, str]]:
    """经历文本 → 结构化事实卡（主张校验）。"""
    return _engine.enrich(raw_text)


def generate_resume(experiences: List[Dict[str, str]],
                    target_role: str = "",
                    format: str = "markdown",
                    chat_fn: Optional[Callable] = None) -> str:
    """经历 → 定向简历（markdown/html/docx/pdf——docx/pdf 返回文件路径）。"""
    eng = ResumeEngine(chat_fn=chat_fn or _engine.chat_fn)
    if format == "html":
        return eng.to_html(experiences, target_role)
    if format == "docx":
        return eng.to_docx(experiences, target_role)
    if format == "pdf":
        return eng.to_pdf(experiences, target_role)
    return eng.compose(experiences, target_role)


def resume_text_to_docx(resume_text: str, target_role: str = "",
                        out_path: str = "") -> str:
    """对话式收集的简历文本（markdown-ish）→ Word 简历文件。

    前端三栏对话式收集产出的 `_resumeText`（`### 小节` + 内容行）直接渲染为 docx，
    避免 /api/export 只认 experiences 而丢弃对话收集内容（前端契约对齐）。

    Returns: 生成的 .docx 文件路径（python-docx 缺失时降级为 .txt）。
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        import tempfile
        if not out_path:
            out_path = os.path.join(tempfile.gettempdir(), "resume_plain.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(resume_text or "")
        return out_path

    import tempfile
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题（居中）
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(target_role or "个人简历")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)
    doc.add_paragraph()

    for raw in (resume_text or "").split("\n"):
        line = raw.rstrip()
        s = line.strip()
        if not s:
            continue
        if s.startswith("# ") or s.startswith("## ") or s.startswith("### "):
            h = doc.add_paragraph()
            hr = h.add_run(s.lstrip("#").strip())
            hr.font.size = Pt(13)
            hr.bold = True
            hr.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)
        elif s.startswith("**"):
            p = doc.add_paragraph()
            pr = p.add_run(s.replace("**", ""))
            pr.font.size = Pt(10)
            pr.font.color.rgb = RGBColor(0x5C, 0x53, 0x4A)
        else:
            p = doc.add_paragraph()
            p.add_run(s)
            p.paragraph_format.space_after = Pt(4)

    if not out_path:
        out_path = os.path.join(tempfile.gettempdir(), "resume.docx")
    doc.save(out_path)
    return out_path


def list_role_packs() -> List[str]:
    """可用通用 Role Pack 清单。"""
    return sorted(_load_role_packs().keys())
